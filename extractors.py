import streamlit as st
import pytesseract
from PIL import Image
from docx import Document

try:
    from pyzbar.pyzbar import decode as decode_barcode
    BARCODE_AVAILABLE = True
except ImportError:
    BARCODE_AVAILABLE = False


@st.cache_data
def ocr_image_bytes(image_bytes: bytes) -> str:
    img = Image.open(image_bytes)
    return pytesseract.image_to_string(img)


def extract_text_from_image(image_file) -> str:
    """
    image_file is a file-like object from Streamlit.
    We don't cache directly on it, but could if we read bytes.
    """
    img = Image.open(image_file)
    return pytesseract.image_to_string(img)


def extract_barcodes_from_image(image_file):
    if not BARCODE_AVAILABLE:
        return []
    img = Image.open(image_file).convert("L")
    codes = decode_barcode(img)
    values = []
    for c in codes:
        try:
            values.append(c.data.decode("utf-8"))
        except Exception:
            continue
    return values


@st.cache_data
def extract_text_from_docx_bytes(doc_bytes: bytes) -> str:
    from io import BytesIO

    doc = Document(BytesIO(doc_bytes))
    result = []
    for para in doc.paragraphs:
        if para.text.strip():
            result.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) == 2:
                result.append(f"{cells[0]} {cells[1]}")
            elif cells:
                result.append(" | ".join(cells))
    return "\n".join(result)


def extract_text_from_docx(docx_file) -> str:
    # docx_file is a file-like from Streamlit
    return extract_text_from_docx_bytes(docx_file.read())
