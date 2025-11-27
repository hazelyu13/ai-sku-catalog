# ===============================================================
#  TARTE BEAUTY  ·  LUXURY DASHBOARD
#  AI-Powered SKU Intake System
# ===============================================================

import os
import re
import shutil
import platform
import subprocess
import base64
import sqlite3

import streamlit as st
import pytesseract
from PIL import Image
from docx import Document
import pandas as pd
import openpyxl

# Try to import barcode / QR decoder
try:
    from pyzbar.pyzbar import decode as decode_barcode
    BARCODE_AVAILABLE = True
except ImportError:
    BARCODE_AVAILABLE = False

# ===============================================================
# GLOBAL CONFIG
# ===============================================================
EXCEL_PATH = "data/specs/sample_specs.xlsx"
EXCEL_SHEET_NAME = "Master Sheet - 12th Floor"
EXCEL_INSERT_ROW = 757

st.set_page_config(
    page_title="Tarte AI SKU System",
    page_icon="💜",
    layout="wide",
)

# ===============================================================
# LUXURY TARTE UI THEME (FULL CSS)
# ===============================================================
tarte_css = """
<style>

    /* GLOBAL BACKGROUND - solid tarte purple */
    .stApp {
        background: #c098ea !important;
        background-attachment: fixed;
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    }

    /* MAIN CONTENT WIDTH */
    div.block-container {
        padding-top: 0 !important;
        margin-top: 0 !important;
        background: transparent !important;
        box-shadow: none !important;
        border: none !important;
    }

    /* TOP TOOLBAR */
    header[data-testid="stHeader"] {
        background-color: rgba(0, 0, 0, 0);
    }

    /* SIDEBAR PANEL */
    section[data-testid="stSidebar"] {
        background: #e1daf8;
        border-right: 2px solid #8b55c9;
        padding-top: 2rem;
    }

    /* Remove radio bullets */
    section[data-testid="stSidebar"] div[role="radiogroup"] > label > div:first-child {
        display: none !important;
    }

    /* Style menu items */
    section[data-testid="stSidebar"] div[role="radiogroup"] > label {
        padding: 10px 14px;
        border-radius: 12px;
        cursor: pointer;
        font-size: 17px;
        font-weight: 600;
        color: #2A0F44 !important;
    }

    /* Hover effect */
    section[data-testid="stSidebar"] div[role="radiogroup"] > label:hover {
        background: #FFFFFF !important;
        color: #2A0F44 !important;
    }

    /* Selected item */
    section[data-testid="stSidebar"] div[role="radiogroup"] [aria-checked="true"] {
        background: rgba(255, 255, 255, 0.5) !important;
        border-radius: 12px;
    }

    section[data-testid="stSidebar"] * {
        color: #240a3f !important;
        font-weight: 500;
    }

    /* Sidebar title */
    section[data-testid="stSidebar"] h1,
    section[data-testid="stSidebar"] h2,
    section[data-testid="stSidebar"] h3 {
        color: #240a3f !important;
        font-weight: 800 !important;
    }

    /* Sidebar radio as pill nav */
    div[role="radiogroup"] > label {
        border-radius: 999px;
        padding: 0.3rem 0.9rem;
        margin-bottom: 0.25rem;
        transition: all 0.18s ease-in-out;
        border: 1px solid transparent;
    }

    div[role="radiogroup"] > label:hover {
        background-color: #f6edff40;
        border-color: #f6edff80;
    }

    /* Sidebar Open Excel button - white version */
    section[data-testid="stSidebar"] div.stButton > button {
        width: 100%;
        background: #ffffff !important;
        color: #2A0F44 !important;
        border-radius: 999px !important;
        border: 2px solid #d0c2e8 !important;
        font-weight: 700 !important;
        padding: 0.6rem 1.2rem;
        font-size: 15px !important;
        box-shadow: 0 4px 12px rgba(0,0,0,0.08);
    }
    section[data-testid="stSidebar"] div.stButton > button:hover {
        background: #f4efff !important;
        border-color: #b9a6dd !important;
        transform: translateY(-1px);
    }

    /* LUX HEADER BANNER */
    .tarte-header {
        background: linear-gradient(90deg, #ffffffee, #f7efffff);
        padding: 26px 40px;
        border-radius: 26px;
        border: 1px solid #ebddff;
        box-shadow: 0 12px 30px rgba(115, 71, 153, 0.20);
        display: flex;
        align-items: center;
        justify-content: space-between;
        gap: 1.5rem;
        margin-bottom: 24px;
    }

    .tarte-header-left {
        text-align: left;
    }

    .tarte-header-title {
        font-size: 30px;
        font-weight: 800;
        color: #34114f;
        margin: 0;
    }

    .tarte-header-sub {
        margin: 4px 0 0 0;
        font-size: 14px;
        color: #62447f;
    }

    .tarte-header-pill {
        background: #f5ecff;
        border-radius: 999px;
        padding: 8px 16px;
        font-size: 13px;
        color: #6b3ea5;
        border: 1px solid #ecd9ff;
        display: inline-flex;
        align-items: center;
        gap: 6px;
    }

    /* SECTION TITLES */
    h1, h2, h3, h4, h5 {
        color: #34114f !important;
        font-weight: 750 !important;
    }

    /* APPLY BLACK TEXT ONLY TO MAIN PAGE (not sidebar) */
    div.block-container li,
    div.block-container h1,
    div.block-container h2,
    div.block-container h3,
    div.block-container h4,
    div.block-container h5 {
        color: #000000 !important;
    }

    /* Only normal body text should be black */
    div.block-container p,
    div.block-container li {
        color: #000000 !important;
    }

    /* Headers keep the theme color */
    div.block-container h1,
    div.block-container h2,
    div.block-container h3 {
        color: #34114f !important;
    }



    /* CHAT BUBBLES */
    .stChatMessage {
        background: #ffffffdd !important;
        border-radius: 12px !important;
        padding: 12px !important;
    }

    /* MAIN PAGE BUTTONS (not sidebar) */
    div.block-container div[data-testid="stButton"] > button {
        background: #ffffff !important;
        color: #2A0F44 !important;
        border: 2px solid #d0c2e8 !important;
        border-radius: 12px !important;
        padding: 8px 16px !important;
        font-weight: 700 !important;
        font-size: 15px !important;
        box-shadow: 0 4px 10px rgba(0,0,0,0.08);
    }

    div.block-container div[data-testid="stButton"] > button:hover {
        background: #f5efff !important;
        border-color: #b9a6dd !important;
        transform: translateY(-1px);
    }

    /* SELECTBOX / DROPDOWN CUSTOMIZATION */
    /* Closed state */
    div[data-baseweb="select"] > div {
        background: #ffffff !important;
        border-radius: 12px !important;
        border: 2px solid #d0c2e8 !important;
        color: #2A0F44 !important;
        font-weight: 500 !important;
        min-height: 44px !important;
        padding: 0 12px !important;
        display: flex !important;
        align-items: center !important;
    }

    div[data-baseweb="select"] span {
        font-size: 14px !important;
        line-height: 1.2 !important;
    }

    /* DROPDOWN MENU PANEL */
    div[data-baseweb="select"] ~ div {
        background: #ffffff !important;
        border-radius: 12px !important;
        border: 2px solid #d0c2e8 !important;
        box-shadow: 0 8px 20px rgba(0,0,0,0.15) !important;
        padding: 6px 0 !important;
    }

    /* DROPDOWN ITEM STYLING */
    div[data-baseweb="select"] ~ div > div {
        background: #ffffff !important;
        color: #ffffff !important;
        font-weight: 600 !important;
        padding: 10px 14px !important;
    }

    /* HOVER */
    div[data-baseweb="select"] ~ div > div:hover {
        background: #f5e9ff !important;
        color: #34114f !important;
    }

    /* ONLY FIX DROPDOWN TEXT COLOR */
    div[data-baseweb="select"] ~ div * {
        color: #ffffff !important;
        font-weight: 600 !important;
    }

    /* Make only the uploader box text white */
[data-testid="stFileUploadDropzone"] * {
    color: #ffffff !important;
    font-weight: 600 !important;
}


</style>
"""
st.markdown(tarte_css, unsafe_allow_html=True)

# ===============================================================
# Helpers
# ===============================================================


def open_excel_file(excel_path: str = EXCEL_PATH) -> None:
    """Open the Excel file using the OS default program."""
    try:
        system = platform.system()
        if system == "Darwin":
            subprocess.call(["open", excel_path])
        elif system == "Windows":
            os.startfile(excel_path)  # type: ignore[attr-defined]
        else:
            subprocess.call(["xdg-open", excel_path])
        st.success("📂 Excel file opened!")
    except Exception as e:
        st.error(f"❌ Could not open Excel file.\n\n{e}")


def load_image_base64(path: str) -> str:
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()


# ===============================================================
# 1. OCR / BARCODE / DOCX EXTRACTION
# ===============================================================


def extract_text_from_image(image_file) -> str:
    img = Image.open(image_file)
    return pytesseract.image_to_string(img)


def extract_barcodes_from_image(image_file):
    if not BARCODE_AVAILABLE:
        return []
    img = Image.open(image_file).convert("L")
    codes = decode_barcode(img)
    values = []
    for c in codes:
        try:
            values.append(c.data.decode("utf-8"))
        except Exception:
            continue
    return values


def extract_text_from_docx(docx_file) -> str:
    doc = Document(docx_file)
    result = []

    for para in doc.paragraphs:
        if para.text.strip():
            result.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if len(cells) == 2:
                result.append(f"{cells[0]} {cells[1]}")
            elif cells:
                result.append(" | ".join(cells))

    return "\n".join(result)


# ===============================================================
# 2. PARSE FIELDS
# ===============================================================


def parse_vendor_doc(text: str) -> dict:
    fields = {
        "Product Description": None,
        "Batch/Lot No.": None,
        "Date": None,
        "SKU": None,
        "Qty": None,
    }
    for line in text.splitlines():
        line = line.strip()
        for key in fields.keys():
            if line.lower().startswith(key.lower()):
                parts = line.split(":", 1)
                if len(parts) == 2:
                    fields[key] = parts[1].strip()
    return fields


# ===============================================================
# 2B. AI FIELD PREDICTION (returns fields + ai_flags)
# ===============================================================


def predict_missing_fields(raw_text: str, fields: dict):
    """
    Enhance vendor field extraction using regex + heuristics.
    Returns (fields, ai_flags) where ai_flags marks which fields AI filled.
    """
    ai_flags = {k: False for k in fields.keys()}
    text = raw_text.lower()

    # SKU
    if not fields.get("SKU"):
        sku_match = re.search(
            r"(sku|style|item)[^\w]?[\s#:]*([a-z0-9\-]{3,20})",
            text,
        )
        if sku_match:
            fields["SKU"] = sku_match.group(2).upper()
            ai_flags["SKU"] = True

    # Lot / Batch
    if not fields.get("Batch/Lot No."):
        lot_match = re.search(
            r"(lot|batch|code|lote)[^\w]?[\s#:]*([a-z0-9\-]{3,20})",
            text,
        )
        if lot_match:
            fields["Batch/Lot No."] = lot_match.group(2).upper()
            ai_flags["Batch/Lot No."] = True

    # Date
    if not fields.get("Date"):
        date_match = re.search(
            r"((\d{4}[-./]\d{1,2}[-./]\d{1,2})|(\d{1,2}[-./]\d{1,2}[-./]\d{2,4}))",
            text,
        )
        if date_match:
            fields["Date"] = date_match.group(1)
            ai_flags["Date"] = True

    # Qty
    if not fields.get("Qty"):
        qty_match = re.search(
            r"(qty|quantity|pcs|units|pack)[^\d]*(\d{1,5})",
            text,
        )
        if qty_match:
            fields["Qty"] = qty_match.group(2)
            ai_flags["Qty"] = True

    # Product description
    if not fields.get("Product Description"):
        desc_match = re.search(
            r"(lipstick|concealer|foundation|palette|mascara|gel|serum|cream|gloss)[a-z0-9\s\-]*",
            text,
        )
        if desc_match:
            fields["Product Description"] = desc_match.group(0).title()
            ai_flags["Product Description"] = True
        else:
            for line in raw_text.splitlines():
                line_stripped = line.strip()
                if 20 <= len(line_stripped) <= 80:
                    fields["Product Description"] = line_stripped
                    ai_flags["Product Description"] = True
                    break

    return fields, ai_flags


# ===============================================================
# 2C. REVIEW UI (with AI-filled chip support)
# ===============================================================


def review_fields_ui(initial_fields, raw_text, key_prefix="", ai_flags=None):
    """
    Side-by-side review:
    Left: raw extracted text
    Right: editable fields with AI-filled badges where applicable.
    """
    if ai_flags is None:
        ai_flags = {}

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📄 Extracted Text")
        st.text_area(
            "Raw OCR / DOCX text",
            value=raw_text,
            height=260,
            key=f"{key_prefix}_raw_text",
        )

    edited = {}

    def field_row(label, field_key, default, widget_key):
        st.markdown(
            f"""
            <div style="display:flex; align-items:center; justify-content:space-between; margin-bottom:2px;">
                <span style="font-weight:600; color:#34114f;">{label}</span>
                {"<span style='background:linear-gradient(90deg,#f5ecff,#e4d4ff); padding:2px 10px; border-radius:999px; font-size:11px; color:#6b3ea5; border:1px solid #e0c8ff;'>✨ AI-filled</span>" if ai_flags.get(field_key) else ""}
            </div>
            """,
            unsafe_allow_html=True,
        )
        return st.text_input(
            label="",
            value=default or "",
            key=widget_key,
        )

    with col2:
        st.subheader("✏️ Review & Edit Fields")

        edited["Product Description"] = field_row(
            "Product Description",
            "Product Description",
            initial_fields.get("Product Description"),
            f"{key_prefix}_desc",
        )
        edited["Batch/Lot No."] = field_row(
            "Batch / Lot No.",
            "Batch/Lot No.",
            initial_fields.get("Batch/Lot No."),
            f"{key_prefix}_lot",
        )
        edited["Date"] = field_row(
            "Date",
            "Date",
            initial_fields.get("Date"),
            f"{key_prefix}_date",
        )
        edited["SKU"] = field_row(
            "SKU",
            "SKU",
            initial_fields.get("SKU"),
            f"{key_prefix}_sku",
        )
        edited["Qty"] = field_row(
            "Qty",
            "Qty",
            initial_fields.get("Qty"),
            f"{key_prefix}_qty",
        )

    return edited


# ===============================================================
# 3. DATABASE + DUPLICATE CHECK
# ===============================================================


def init_db():
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS sku_catalog (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            product_desc TEXT,
            batch_lot TEXT,
            date TEXT,
            sku TEXT,
            qty TEXT
        )
        """
    )
    conn.close()


def save_to_db(fields):
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute(
        """
        INSERT INTO sku_catalog (product_desc, batch_lot, date, sku, qty)
        VALUES (?, ?, ?, ?, ?)
        """,
        (
            fields.get("Product Description"),
            fields.get("Batch/Lot No."),
            fields.get("Date"),
            fields.get("SKU"),
            fields.get("Qty"),
        ),
    )
    conn.commit()
    conn.close()
    st.success("✅ Saved to database!")


def check_duplicate(fields):
    sku = fields.get("SKU")
    lot = fields.get("Batch/Lot No.")
    if not sku or not lot:
        return None

    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute(
        """
        SELECT * FROM sku_catalog
        WHERE sku = ? AND batch_lot = ?
        """,
        (sku, lot),
    )
    result = cur.fetchone()
    conn.close()
    return result


def search_db(keyword):
    conn = sqlite3.connect("sku_catalog.db")
    df = pd.read_sql_query(
        f"""
        SELECT * FROM sku_catalog
        WHERE sku LIKE '%{keyword}%'
        OR batch_lot LIKE '%{keyword}%'
        OR product_desc LIKE '%{keyword}%'
        """,
        conn,
    )
    conn.close()
    return df


def clear_database():
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM sku_catalog")
    conn.commit()
    conn.close()


# ===============================================================
# 3B. EXTRA DB HELPERS (Chatbot)
# ===============================================================


def export_csv():
    conn = sqlite3.connect("sku_catalog.db")
    df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
    conn.close()
    export_path = "sku_export.csv"
    df.to_csv(export_path, index=False)
    return export_path


def export_excel():
    conn = sqlite3.connect("sku_catalog.db")
    df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
    conn.close()
    export_path = "sku_export.xlsx"
    df.to_excel(export_path, index=False)
    return export_path


def clear_row(row_id):
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute("DELETE FROM sku_catalog WHERE id = ?", (row_id,))
    conn.commit()
    conn.close()


def undo_last_save():
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute("SELECT id FROM sku_catalog ORDER BY id DESC LIMIT 1")
    last = cur.fetchone()
    if last:
        cur.execute("DELETE FROM sku_catalog WHERE id = ?", (last[0],))
        conn.commit()
        conn.close()
        return last[0]
    conn.close()
    return None


def database_stats():
    conn = sqlite3.connect("sku_catalog.db")
    cur = conn.cursor()
    cur.execute("SELECT COUNT(*) FROM sku_catalog")
    count = cur.fetchone()[0]
    cur.execute(
        """
        SELECT sku, batch_lot, COUNT(*)
        FROM sku_catalog
        GROUP BY sku, batch_lot
        HAVING COUNT(*) > 1
        """
    )
    dups = cur.fetchall()
    conn.close()
    return count, dups


# ===============================================================
# 4. SAVE TO EXCEL
# ===============================================================


def save_to_excel(fields):
    if not os.path.exists(EXCEL_PATH):
        st.error(f"❌ Excel not found at {EXCEL_PATH}")
        return

    temp_path = EXCEL_PATH.replace(".xlsx", "_temp.xlsx")
    shutil.copyfile(EXCEL_PATH, temp_path)

    wb = openpyxl.load_workbook(temp_path)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        st.error("❌ Sheet name incorrect.")
        return

    ws = wb[EXCEL_SHEET_NAME]

    excel_headers = {}
    for col in range(1, ws.max_column + 1):
        header_value = ws.cell(row=1, column=col).value
        if header_value:
            excel_headers[header_value.lower()] = col

    mapping = {
        "product description": fields.get("Product Description"),
        "lot #": fields.get("Batch/Lot No."),
        "date": fields.get("Date"),
        "sku": fields.get("SKU"),
        "qty": fields.get("Qty"),
    }

    target_row = EXCEL_INSERT_ROW
    while ws.max_row < target_row:
        ws.append([])

    for key, value in mapping.items():
        if key in excel_headers:
            col = excel_headers[key]
            ws.cell(row=target_row, column=col, value=value)

    wb.save(temp_path)
    shutil.move(temp_path, EXCEL_PATH)
    st.success("💾 Saved to Excel successfully!")


# ===============================================================
# 5. GALLERY IMAGE SAVER (SKU + LOT in filename)
# ===============================================================


def save_image_to_gallery(fields, file_obj, original_name):
    gallery_folder = "data/gallery"
    os.makedirs(gallery_folder, exist_ok=True)

    sku = fields.get("SKU") or "NOSKU"
    lot = fields.get("Batch/Lot No.") or "NOLOT"

    sku_safe = re.sub(r"[^A-Za-z0-9\-]", "", str(sku))
    lot_safe = re.sub(r"[^A-Za-z0-9\-]", "", str(lot))

    ext = os.path.splitext(original_name)[1].lower()
    if ext not in [".jpg", ".jpeg", ".png"]:
        ext = ".jpg"

    filename = f"SKU_{sku_safe}_LOT_{lot_safe}{ext}"
    path = os.path.join(gallery_folder, filename)

    with open(path, "wb") as out:
        out.write(file_obj.getbuffer())

    return path


# Ensure DB exists
init_db()

# ===============================================================
# SIDEBAR NAV + LOGO
# ===============================================================
logo_base64 = None
try:
    logo_base64 = load_image_base64("data/images/tarte_logo.png")
except FileNotFoundError:
    logo_base64 = None

if logo_base64:
    st.sidebar.markdown(
        f"""
        <div style="text-align:center; margin-bottom:10px;">
            <img src="data:image/png;base64,{logo_base64}"
                 style="width:140px; margin-top:-10px;">
        </div>
        """,
        unsafe_allow_html=True,
    )

st.sidebar.title("Tarte SKU System")

page = st.sidebar.radio(
    "",
    ["Home", "Upload", "Camera", "Chatbot", "Database", "Gallery", "Excel Tools"],
    label_visibility="collapsed",
)

if st.sidebar.button("📂 Open Excel File"):
    open_excel_file()

# ===============================================================
# HEADER BANNER (SHARED)
# ===============================================================
st.markdown(
    """
    <div class="tarte-header">
        <div class="tarte-header-left">
            <p class="tarte-header-title">Tarte AI-Powered SKU Intake</p>
            <p class="tarte-header-sub">
                Designed for production retains, lab samples, and go to market traceability.
            </p>
        </div>
        <div>
            <span class="tarte-header-pill">
                🧾 Auto intake · ✅ Duplicate safe · 📊 Excel synced
            </span>
        </div>
    </div>
    """,
    unsafe_allow_html=True,
)

# ===============================================================
# HOME PAGE
# ===============================================================
if page == "Home":
    st.subheader("✨ Welcome")
    st.write(
        """
This dashboard automates the workflow for cataloging Tarte production retains:

• Upload DOCX or image files  
• Capture new retains with the camera  
• Extract text + barcodes and parse core SKU fields  
• Review and edit before saving  
• Store records in a duplicate-safe SQLite database  
• Sync rows directly into the master Excel tracker  
• Use the chatbot to trigger actions or search data  
        """
    )

# ===============================================================
# UPLOAD PAGE
# ===============================================================
elif page == "Upload":

    st.header("📄 Upload Vendor Documents")
    st.write("Upload one or more DOCX or image files. Review fields before saving.")

    files = st.file_uploader(
        "Upload one or more files",
        type=["docx", "jpg", "jpeg", "png"],
        accept_multiple_files=True,
        help="Upload retain specs, artwork proofs, or vendor docs.",
    )

    if files:
        st.success(f"{len(files)} files uploaded.")

        for idx, f in enumerate(files, start=1):
            st.markdown("---")
            st.subheader(f"📄 File {idx}: {f.name}")

            if f.name.lower().endswith(".docx"):
                text = extract_text_from_docx(f)
                barcodes = []
                is_image_file = False
            else:
                text = extract_text_from_image(f)
                barcodes = extract_barcodes_from_image(f)
                is_image_file = True

            if barcodes:
                st.info(f"🔍 Detected barcodes / QR codes: {', '.join(barcodes)}")

            fields = parse_vendor_doc(text)
            fields, ai_flags = predict_missing_fields(text, fields)

            if barcodes and not fields.get("SKU"):
                fields["SKU"] = barcodes[0]

            duplicate_row = check_duplicate(fields)
            if duplicate_row:
                st.warning(
                    f"⚠️ Duplicate detected for SKU **{fields.get('SKU')}** / "
                    f"Lot **{fields.get('Batch/Lot No.')}**",
                    icon="🚨",
                )
                dup_df = pd.DataFrame(
                    [duplicate_row],
                    columns=["id", "product_desc", "batch_lot", "date", "sku", "qty"],
                )
                st.dataframe(dup_df, use_container_width=True)

            edited_fields = review_fields_ui(
                fields,
                text,
                key_prefix=f"file_{idx}",
                ai_flags=ai_flags,
            )

            col_a, col_b = st.columns(2)
            with col_a:
                confirm = st.button(
                    f"✅ Confirm & Save {f.name}",
                    key=f"confirm_{idx}",
                )
            with col_b:
                cancel = st.button(
                    f"❌ Cancel {f.name}",
                    key=f"cancel_{idx}",
                )

            if confirm:
                duplicate_row_final = check_duplicate(edited_fields)
                if duplicate_row_final:
                    st.error("🚫 Duplicate detected - saving blocked.")
                    if st.button(
                        "🔓 Override & Force Save",
                        key=f"override_{idx}",
                    ):
                        save_to_db(edited_fields)
                        save_to_excel(edited_fields)
                        if is_image_file:
                            save_image_to_gallery(edited_fields, f, f.name)
                else:
                    save_to_db(edited_fields)
                    save_to_excel(edited_fields)
                    if is_image_file:
                        save_image_to_gallery(edited_fields, f, f.name)

            elif cancel:
                st.info(f"⏹ Skipped saving for {f.name}.")

# ===============================================================
# CAMERA PAGE
# ===============================================================
elif page == "Camera":

    st.header("📸 Capture Vendor Document")
    photo = st.camera_input("Take a photo of the retain or vendor sheet")

    if photo:
        text = extract_text_from_image(photo)
        barcodes = extract_barcodes_from_image(photo)

        if barcodes:
            st.info(f"🔍 Detected barcodes / QR codes: {', '.join(barcodes)}")

        fields = parse_vendor_doc(text)
        fields, ai_flags = predict_missing_fields(text, fields)

        if barcodes and not fields.get("SKU"):
            fields["SKU"] = barcodes[0]

        duplicate_row = check_duplicate(fields)
        if duplicate_row:
            st.warning(
                f"⚠️ Duplicate detected for SKU **{fields.get('SKU')}** / "
                f"Lot **{fields.get('Batch/Lot No.')}**",
                icon="🚨",
            )
            dup_df = pd.DataFrame(
                [duplicate_row],
                columns=["id", "product_desc", "batch_lot", "date", "sku", "qty"],
            )
            st.dataframe(dup_df, use_container_width=True)

        edited_fields = review_fields_ui(
            fields,
            text,
            key_prefix="camera",
            ai_flags=ai_flags,
        )

        col_a, col_b = st.columns(2)
        with col_a:
            confirm = st.button("✅ Confirm & Save from camera")
        with col_b:
            cancel = st.button("❌ Cancel (discard photo)")

        if confirm:
            duplicate_row_final = check_duplicate(edited_fields)
            if duplicate_row_final:
                st.error("🚫 Duplicate detected - saving blocked.")
                if st.button("🔓 Override & Force Save", key="override_camera"):
                    save_to_db(edited_fields)
                    save_to_excel(edited_fields)
                    save_image_to_gallery(edited_fields, photo, "camera.jpg")
            else:
                save_to_db(edited_fields)
                save_to_excel(edited_fields)
                save_image_to_gallery(edited_fields, photo, "camera.jpg")

        elif cancel:
            st.info("⏹ Camera capture discarded.")

# ===============================================================
# CHATBOT PAGE
# ===============================================================
elif page == "Chatbot":

    st.header("🤖 Chatbot Mode")
    st.write(
        "Try commands like: `process latest`, `search SN52`, `show database`, "
        "`clear database`, `export csv`, `export excel`, `undo`, `stats`, "
        "`list duplicates`, `open excel`, `show latest`, `help`."
    )

    if "messages" not in st.session_state:
        st.session_state.messages = []

    for m in st.session_state.messages:
        with st.chat_message(m["role"]):
            st.write(m["content"])

    user_input = st.chat_input("Ask me to process latest, search, or run commands")

    if user_input:
        st.session_state.messages.append({"role": "user", "content": user_input})
        reply = ""
        lower = user_input.lower()

        if "process latest" in lower:
            folder = "data/specs/"
            files = [
                os.path.join(folder, f)
                for f in os.listdir(folder)
                if f.lower().endswith((".docx", ".jpg", ".jpeg", ".png"))
            ]

            if not files:
                reply = "❌ No files found in data/specs."
            else:
                latest = max(files, key=os.path.getmtime)

                if latest.lower().endswith(".docx"):
                    with open(latest, "rb") as fh:
                        text = extract_text_from_docx(fh)
                    barcodes = []
                    is_image_file = False
                else:
                    with open(latest, "rb") as fh:
                        text = extract_text_from_image(fh)
                        fh.seek(0)
                        barcodes = extract_barcodes_from_image(fh)
                    is_image_file = True

                fields = parse_vendor_doc(text)
                fields, _ = predict_missing_fields(text, fields)

                if barcodes and not fields.get("SKU"):
                    fields["SKU"] = barcodes[0]

                duplicate_row = check_duplicate(fields)
                if duplicate_row:
                    reply = (
                        f"⚠️ Duplicate detected for SKU {fields.get('SKU')} / "
                        f"Lot {fields.get('Batch/Lot No.')} - not saved.\n\n"
                        "Use the Upload or Camera page to review and override if needed."
                    )
                else:
                    save_to_db(fields)
                    save_to_excel(fields)
                    reply = (
                        f"Processed and saved latest file: "
                        f"{os.path.basename(latest)}"
                    )

                    if is_image_file:
                        with open(latest, "rb") as fh:
                            class _Tmp:
                                def __init__(self, b):
                                    self._b = b

                                def getbuffer(self):
                                    return self._b

                            buf = fh.read()
                            tmp_file = _Tmp(buf)
                            save_image_to_gallery(
                                fields, tmp_file, os.path.basename(latest)
                            )

        elif "search" in lower:
            term = user_input.split("search", 1)[1].strip()
            df = search_db(term)
            st.dataframe(df)
            reply = f"Search results for '{term}'."

        elif "show database" in lower:
            conn = sqlite3.connect("sku_catalog.db")
            df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
            conn.close()
            st.dataframe(df)
            reply = "Full database view."

        elif "clear database" in lower:
            clear_database()
            reply = "Database cleared."

        elif "export csv" in lower:
            path = export_csv()
            reply = f"📤 Exported as CSV: `{path}`"

        elif "export excel" in lower:
            path = export_excel()
            reply = f"📤 Exported as Excel: `{path}`"

        elif "undo" in lower:
            deleted = undo_last_save()
            if deleted:
                reply = f"↩️ Undid last save (deleted row ID {deleted})."
            else:
                reply = "Nothing to undo."

        elif "stats" in lower:
            total, duplicates = database_stats()
            reply = f"📊 Total rows: {total}\n\nDuplicate SKU/Lot combos:\n{duplicates}"

        elif "delete" in lower:
            term = lower.split("delete", 1)[1].strip()
            conn = sqlite3.connect("sku_catalog.db")
            cur = conn.cursor()
            cur.execute("DELETE FROM sku_catalog WHERE sku = ?", (term,))
            conn.commit()
            conn.close()
            reply = f"🗑 Deleted all rows with SKU `{term}`."

        elif "list duplicates" in lower:
            conn = sqlite3.connect("sku_catalog.db")
            df = pd.read_sql_query(
                """
                SELECT sku, batch_lot, COUNT(*) as count
                FROM sku_catalog
                GROUP BY sku, batch_lot
                HAVING COUNT(*) > 1
                """,
                conn,
            )
            conn.close()
            st.dataframe(df)
            reply = "🔍 Duplicate rows shown above."

        elif "open excel" in lower:
            open_excel_file()
            reply = "Opening Excel..."

        elif "show latest" in lower:
            conn = sqlite3.connect("sku_catalog.db")
            df = pd.read_sql_query(
                "SELECT * FROM sku_catalog ORDER BY id DESC LIMIT 1",
                conn,
            )
            conn.close()
            st.dataframe(df)
            reply = "📄 Latest saved row shown above."

        elif "help" in lower:
            reply = """
Available Commands:
• process latest  
• search <term>  
• show database  
• clear database  
• export csv  
• export excel  
• delete <SKU>  
• undo  
• stats  
• list duplicates  
• open excel  
• show latest  
• help
"""

        else:
            reply = (
                "Unknown command. Try: `process latest`, `search <term>`, "
                "`show database`, `clear database`, `export csv`, `export excel`, "
                "`undo`, `stats`, `list duplicates`, `open excel`, `show latest`, or `help`."
            )

        st.session_state.messages.append({"role": "assistant", "content": reply})

# ===============================================================
# DATABASE PAGE
# ===============================================================
elif page == "Database":

    st.header("📊 Database Viewer")

    search = st.text_input("", placeholder="🔍 Search SKU, lot, or product description")

    if search:
        df = search_db(search)
        st.dataframe(df)
    else:
        conn = sqlite3.connect("sku_catalog.db")
        df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
        conn.close()
        st.dataframe(df)

    if st.button("🧹 Clear database"):
        clear_database()
        st.success("Database cleared.")

# ===============================================================
# GALLERY PAGE (Pinterest-style with filters)
# ===============================================================
elif page == "Gallery":

    st.header("🖼️ Processed Retains Gallery")
    st.write("A visual grid of all saved retains, with filters for SKU, lot, and date.")

    gallery_folder = "data/gallery"

    if not os.path.exists(gallery_folder):
        os.makedirs(gallery_folder)

    images = [
        os.path.join(gallery_folder, img)
        for img in os.listdir(gallery_folder)
        if img.lower().endswith((".jpg", ".jpeg", ".png"))
    ]

    conn = sqlite3.connect("sku_catalog.db")
    db_df = pd.read_sql_query("SELECT * FROM sku_catalog", conn)
    conn.close()

    if db_df.empty:
        sku_options = []
        lot_options = []
        date_options = []
    else:
        sku_options = sorted(db_df["sku"].dropna().astype(str).unique().tolist())
        lot_options = sorted(db_df["batch_lot"].dropna().astype(str).unique().tolist())
        date_options = sorted(db_df["date"].dropna().astype(str).unique().tolist())

    filt_col1, filt_col2, filt_col3, filt_col4 = st.columns([1, 1, 1, 2])

    with filt_col1:
        sku_filter = st.selectbox(
            "Filter by SKU",
            ["All"] + sku_options if sku_options else ["All"],
        )

    with filt_col2:
        lot_filter = st.selectbox(
            "Filter by Lot",
            ["All"] + lot_options if lot_options else ["All"],
        )

    with filt_col3:
        date_filter = st.selectbox(
            "Filter by Date",
            ["All"] + date_options if date_options else ["All"],
        )

    with filt_col4:
        search_filter = st.text_input(
            "Search text (SKU, lot, or description)",
            "",
            placeholder="e.g. SN52, holiday set, batch 3392",
        ).strip().lower()

    st.markdown("---")

    if not images:
        st.info("No retains in the gallery yet - save from Upload or Camera first.")
    else:
        cols = st.columns(4)

        lookup = {}
        if not db_df.empty:
            for _, row in db_df.iterrows():
                key = (str(row["sku"]), str(row["batch_lot"]))
                lookup[key] = row

        for idx, img_path in enumerate(sorted(images)):
            basename = os.path.basename(img_path)
            match = re.match(
                r"SKU_(.+)_LOT_(.+)\.[^.]+$",
                basename,
                re.IGNORECASE,
            )

            display_info = None
            sku_val = None
            lot_val = None

            if match:
                sku_val = match.group(1)
                lot_val = match.group(2)
                display_info = lookup.get((sku_val, lot_val))

            if sku_filter != "All":
                if not sku_val or sku_val != sku_filter:
                    continue

            if lot_filter != "All":
                if not lot_val or lot_val != lot_filter:
                    continue

            if date_filter != "All":
                if display_info is None or str(display_info.get("date")) != date_filter:
                    continue

            if search_filter:
                if display_info is None:
                    continue
                haystack = " ".join(
                    [
                        str(display_info.get("sku", "")).lower(),
                        str(display_info.get("batch_lot", "")).lower(),
                        str(display_info.get("product_desc", "")).lower(),
                    ]
                )
                if search_filter not in haystack:
                    continue

            col = cols[idx % 4]
            with col:
                st.markdown(
                    """
                    <div style="
                        background: #ffffffaa;
                        border-radius: 16px;
                        padding: 12px;
                        margin-bottom: 20px;
                        box-shadow: 0 6px 15px rgba(0,0,0,0.1);
                        backdrop-filter: blur(6px);
                    ">
                    """,
                    unsafe_allow_html=True,
                )

                st.image(img_path, use_column_width=True)

                if display_info is not None:
                    st.markdown(
                        f"""
                        <p style="font-size:14px; margin-top:8px;">
                        <b>SKU:</b> {display_info['sku']}<br>
                        <b>Lot:</b> {display_info['batch_lot']}<br>
                        <b>Date:</b> {display_info['date']}<br>
                        <b>Qty:</b> {display_info['qty']}
                        </p>
                        """,
                        unsafe_allow_html=True,
                    )
                else:
                    st.caption("No matching record found for this image.")

                st.markdown("</div>", unsafe_allow_html=True)

# ===============================================================
# EXCEL TOOLS
# ===============================================================
elif page == "Excel Tools":

    st.header("📂 Excel Tools")
    st.write(f"Excel file: `{EXCEL_PATH}`")
    st.write(f"Sheet name: `{EXCEL_SHEET_NAME}`")
    st.write(f"Inserting into row: `{EXCEL_INSERT_ROW}`")

    if st.button("📂 Open Excel file"):
        open_excel_file()

    st.markdown("---")
    st.write("Tip: close the workbook before running updates from this dashboard.")
